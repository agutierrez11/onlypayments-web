import os
import csv
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import win32com.client

# Target output directories
OUTPUT_DIRS = [
    r"C:\Users\Antonio\Downloads",
    r"C:\Users\Antonio\OneDrive\Escritorio",
]

# Ensure directories exist
for d in OUTPUT_DIRS:
    os.makedirs(d, exist_ok=True)

print("Starting generation of supported formats (DOCX, PDF, TXT, CSV)...")

# ==========================================
# 1. GENERATE DOCX DOSSIER
# ==========================================
doc = Document()

# Set page margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Title
title_p = doc.add_paragraph()
title_run = title_p.add_run("ONLYPAYMENTS — DOSSIER EJECUTIVO 2026")
title_run.font.name = 'Arial'
title_run.font.size = Pt(22)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 150, 140)
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle_p = doc.add_paragraph()
sub_run = subtitle_p.add_run("Arquitectura Técnica, Modelo de Monetización B2B y Dataset de 20 Países en LATAM")
sub_run.font.name = 'Arial'
sub_run.font.size = Pt(12)
sub_run.font.italic = True
sub_run.font.color.rgb = RGBColor(100, 110, 120)
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Heading 1: Visión y Propuesta
h1 = doc.add_heading("1. Visión Ejecutiva y Propuesta de Valor", level=1)
h1.runs[0].font.color.rgb = RGBColor(0, 120, 130)

p1 = doc.add_paragraph()
p1.add_run("OnlyPayments es la Plataforma de Inteligencia Operativa y Marketplace B2B de Introducciones Calificadas para el ecosistema Fintech y Medios de Pago en Latinoamérica.\n\n")
p1.add_run("Propuesta de Valor de Riesgo Cero: ").bold = True
p1.add_run('"Zero cost · Sin mensualidades · Sin contratos obligatorios. El comprador navega gratis; los proveedores pagan únicamente por introducciones B2B calificadas ($150 USD) o comisiones por éxito al cerrar deals de procesamiento."')

# Heading 2: RevOps 70/30
h2 = doc.add_heading("2. Estrategia de Navegación y Conversión (Regla RevOps 70/30)", level=1)
h2.runs[0].font.color.rgb = RGBColor(0, 120, 130)

p2 = doc.add_paragraph("El frontend de OnlyPayments está estructurado bajo una separación psicológica estricta entre la Zona Informativa (70% del scroll) para construir autoridad indiscutible y SEO, y la Zona de Negocio (30% del scroll) para capturar demanda calificada:")

# Table RevOps
table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Zona UI"
hdr_cells[1].text = "% Scroll"
hdr_cells[2].text = "Tono & Copys"
hdr_cells[3].text = "Objetivo Comercial"

for cell in hdr_cells:
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

row1 = table.add_row().cells
row1[0].text = "Zona Informativa"
row1[1].text = "70%"
row1[2].text = '"Te enseñamos el mapa"'
row1[3].text = "Educar, autoridad técnica y SEO orgánico"

row2 = table.add_row().cells
row2[0].text = "Zona de Negocio"
row2[1].text = "30%"
row2[2].text = '"Te ayudamos a negociar"'
row2[3].text = "Captura de Leads (Fintech Matcher) y agendamiento B2B"

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Heading 3: Monetización
h3 = doc.add_heading("3. Motor de Monetización (5 Flujos de Ingreso)", level=1)
h3.runs[0].font.color.rgb = RGBColor(0, 120, 130)

m_list = [
    ("1. Consultoría Gratuita → Comisión (Success Fee): ", "Brokerage de procesamiento B2B. Zero cost para el comprador; cobro de comisión al provider/adquirente al cerrar el contrato."),
    ("2. Sponsored Listings (Listados Destacados): ", "Marcas y pasarelas pagan suscripción mensual ($500 - $2,500 USD) para posicionarse en el top del Mapa 3D y Rankings."),
    ("3. Informes Premium (Market Intelligence): ", "Suscripciones ejecutivas a reportes de regulaciones iGaming, rieles A2A y mercado de remesas."),
    ("4. API de Datos B2B: ", "Venta de API Keys para consultar la base de 1,047+ fintechs y sus stacks de pagos activos."),
    ("5. Matchmaking Pay-per-Intro ($150 USD): ", "Introducción calificada de 15 minutos con decisores (CFOs/CTOs), dispersando 70% al conector, 15% a la comunidad y 15% a la plataforma vía Stripe Connect.")
]

for title, desc in m_list:
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(title)
    r1.bold = True
    p.add_run(desc)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Heading 4: Mapa de Rutas
h4 = doc.add_heading("4. Mapa Oficial de Rutas en la Aplicación Web", level=1)
h4.runs[0].font.color.rgb = RGBColor(0, 120, 130)

routes = [
    ("onlypayments.com/", "Landing Informativa principal (Hero 3D → Radar → Insights → Matcher)"),
    ("onlypayments.com/matcher", "🔥 Fintech Matcher a pantalla completa (Diagnóstico en 4 pasos & Captura)"),
    ("onlypayments.com/ecosistema", "Grafo satelital interactivo completo de 20 países en LATAM"),
    ("onlypayments.com/insights", "Blog y feed de datos del ecosistema (PIX, Bre-B, Remesas)"),
    ("onlypayments.com/rankings", "Tablas sorteables y heatmap de las 50+ fintechs más grandes"),
    ("onlypayments.com/consultoria", "Servicios premium e introducciones B2B ($150 USD/intro)")
]

for url, desc in routes:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(url + " — ")
    r.bold = True
    p.add_run(desc)

# Save DOCX locally first
docx_filename = "DOSSIER_ONLYPAYMENTS_2026.docx"
local_docx_path = os.path.abspath(docx_filename)
doc.save(local_docx_path)
print(f"Saved local DOCX: {local_docx_path}")

# ==========================================
# 2. CONVERT DOCX TO PDF USING MS WORD COM
# ==========================================
pdf_filename = "DOSSIER_ONLYPAYMENTS_2026.pdf"
local_pdf_path = os.path.abspath(pdf_filename)

try:
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc_com = word.Documents.Open(local_docx_path)
    # 17 = wdFormatPDF
    doc_com.SaveAs(local_pdf_path, FileFormat=17)
    doc_com.Close()
    word.Quit()
    print(f"Successfully converted DOCX to PDF: {local_pdf_path}")
except Exception as e:
    print(f"Error converting to PDF via Word COM: {e}")

# ==========================================
# 3. GENERATE TXT DOSSIER
# ==========================================
txt_content = """====================================================================
ONLYPAYMENTS — DOSSIER EJECUTIVO 2026
Arquitectura Técnica, Modelo de Monetización B2B y Dataset de LATAM
====================================================================

1. VISIÓN EJECUTIVA Y PROPUESTA DE VALOR
--------------------------------------------------------------------
OnlyPayments es la Plataforma de Inteligencia Operativa y Marketplace B2B
de Introducciones Calificadas para el ecosistema Fintech y Medios de Pago en LATAM.

Propuesta de Valor: "Zero cost · Sin mensualidades · Sin contratos obligatorios.
El comprador navega gratis; los proveedores pagan únicamente por introducciones B2B
calificadas ($150 USD) o comisiones por éxito al cerrar deals de procesamiento."

2. ESTRATEGIA DE NAVEGACIÓN Y CONVERSIÓN (REVOPS 70/30)
--------------------------------------------------------------------
- Zona Informativa (70% del scroll): Paleta obsidiana (#020408). Tono: "Te enseñamos el mapa".
  Objetivo: Educar, autoridad técnica y SEO orgánico.
- Zona de Negocio (30% del scroll): Acento neón cyan (#00f5d4). Tono: "Te ayudamos a negociar".
  Objetivo: Captura de Leads (Fintech Matcher) y agendamiento B2B.

3. MOTOR DE MONETIZACIÓN (5 FLUJOS DE INGRESO)
--------------------------------------------------------------------
1. Consultoría Gratuita -> Comisión (Success Fee): Brokerage de procesamiento B2B.
2. Sponsored Listings (Listados Destacados): Suscripciones de $500 - $2,500 USD/mes.
3. Informes Premium (Market Intelligence): Reportes de regulaciones y remesas.
4. API de Datos B2B: Acceso a la base de 1,047+ fintechs y stacks.
5. Matchmaking Pay-per-Intro ($150 USD): Split 70% Conector / 15% Comunidad / 15% Plataforma.

4. MAPA OFICIAL DE RUTAS DE LA APLICACIÓN WEB
--------------------------------------------------------------------
- onlypayments.com/            <- Landing Informativa principal (Hero 3D -> Matcher)
- onlypayments.com/matcher     <- Fintech Matcher en 4 pasos (Diagnóstico & Captura Lead)
- onlypayments.com/ecosistema  <- Grafo satelital 3D de 20 países en LATAM
- onlypayments.com/insights    <- Blog de datos del ecosistema (PIX, Bre-B, Remesas)
- onlypayments.com/rankings    <- Tablas sorteables de las 50+ fintechs más grandes
- onlypayments.com/consultoria <- Servicios premium e introducciones B2B ($150 USD)
"""

local_txt_path = os.path.abspath("DOSSIER_ONLYPAYMENTS_2026.txt")
with open(local_txt_path, "w", encoding="utf-8") as f:
    f.write(txt_content)

# ==========================================
# 4. GENERATE CSV DATASET (20 PAÍSES LATAM)
# ==========================================
csv_data = [
    ["País", "Flag", "Fintechs Activas", "VC Funding 2025", "Remesas Anuales", "Riel RTP Dominante", "Regulación / Estatus Crypto"],
    ["México", "🇲🇽", "650+", "$980M USD", "$62.5B USD", "SPEI / DiMo", "Ley Fintech 2018/2026 / VASP"],
    ["Brasil", "🇧🇷", "1,200+", "$2.03B USD", "$1.2B USD", "PIX Automático", "Open Finance Fase 4 / Reg. BCB"],
    ["Colombia", "🇨🇴", "410+", "$92M USD", "$2.1B USD", "Bre-B / PSE", "Decreto Open Banking 2022 / Sandbox"],
    ["Argentina", "🇦🇷", "320+", "Variable", "$0.8B USD", "Transferencias 3.0", "BCRA Interoperable / Adopción USDT"],
    ["Chile", "🇨🇱", "280+", "$110M USD", "$0.3B USD", "TEF / CMF Instant", "Ley Fintech N° 21.521 / Reg. CMF"],
    ["Perú", "🇵🇪", "180+", "Growing", "$4.5B USD", "Yape / Plin", "Interoperabilidad BCRP / SBS"],
    ["Ecuador", "🇪🇨", "90+", "Emerging", "$5.2B USD", "BCE Pago Inmediato", "Ley Fintech 2022 / Dolarización"],
    ["Guatemala", "🇬🇹", "60+", "Seed", "$19.8B USD", "ACH Pronto", "Agenda SIB / Remesas Crypto"],
    ["El Salvador", "🇸🇻", "45+", "Seed", "$8.3B USD", "Bitcoin + Tradicional", "Legal Tender Bitcoin / Chivo"],
    ["Costa Rica", "🇨🇷", "70+", "Seed+", "$0.6B USD", "SINPE Móvil", "Regulación BCCR / Libre Circulación"],
    ["Panamá", "🇵🇦", "85+", "Seed+", "$0.7B USD", "Yappy / ACH Xpress", "Hub Financiero Digital / Dólar"],
    ["Rep. Dominicana", "🇩🇴", "110+", "Series A", "$11.9B USD", "LBTR Inmediato", "Estrategia EBRD"],
    ["Honduras", "🇭🇳", "35+", "Seed", "$9.1B USD", "ACH Banhprovi", "CNBS Sandbox / Próspera ZEDE"],
    ["Uruguay", "🇺🇾", "55+", "Seed+", "$0.15B USD", "SPI / BCU Direct", "Marco BCU 2024"],
    ["Paraguay", "🇵🇾", "40+", "Seed", "$0.4B USD", "SIPAP 24/7", "SIPAP BCP / Minería Crypto"],
    ["Bolivia", "🇧🇴", "30+", "Seed", "$1.4B USD", "QR Simple ASFI", "Normativa ASFI / BCB"],
    ["Nicaragua", "🇳🇮", "25+", "Seed", "$4.2B USD", "ACH UniRed", "Normativa SIBOIF"],
    ["Jamaica", "🇯🇲", "35+", "Seed", "$3.8B USD", "JAM-DEX CBDC", "BOJ CBDC Pioneer"],
    ["Venezuela", "🇻🇪", "50+", "Bootstrap", "$0.1B USD", "Pago Móvil P2P", "Sudeban / Adopción Crypto P2P"],
    ["Puerto Rico", "🇵🇷", "120+", "US Capital", "$1.1B USD", "ATH Móvil", "US Fedwire / Act 60 Fiscal"]
]

local_csv_path = os.path.abspath("LATAM_20_PAISES_DATASET.csv")
with open(local_csv_path, "w", newline="", encoding="utf-8-sig") as f:
    writer = csv.writer(f)
    writer.writerows(csv_data)

print(f"Saved local CSV: {local_csv_path}")

# ==========================================
# 5. GENERATE CODE FILES IN TXT FORMAT
# ==========================================
matcher_html_path = os.path.abspath(r"client/public/fintech-matcher.html")
satellite_html_path = os.path.abspath(r"client/public/latam-fintech-satellite.html")

local_matcher_txt = os.path.abspath("FINTECH_MATCHER_CÓDIGO.txt")
local_satellite_txt = os.path.abspath("SATELLITE_GLOBE_CÓDIGO.txt")

if os.path.exists(matcher_html_path):
    with open(matcher_html_path, "r", encoding="utf-8") as f_in, open(local_matcher_txt, "w", encoding="utf-8") as f_out:
        f_out.write(f_in.read())

if os.path.exists(satellite_html_path):
    with open(satellite_html_path, "r", encoding="utf-8") as f_in, open(local_satellite_txt, "w", encoding="utf-8") as f_out:
        f_out.write(f_in.read())

# ==========================================
# 6. COPY ALL SUPPORTED FILES TO TARGET DIRS
# ==========================================
import shutil

files_to_copy = [
    local_docx_path,
    local_pdf_path if os.path.exists(local_pdf_path) else None,
    local_txt_path,
    local_csv_path,
    local_matcher_txt,
    local_satellite_txt
]

files_to_copy = [f for f in files_to_copy if f and os.path.exists(f)]

for target_dir in OUTPUT_DIRS:
    print(f"\nCopying supported files to: {target_dir}")
    for file_path in files_to_copy:
        fname = os.path.basename(file_path)
        dest = os.path.join(target_dir, fname)
        shutil.copy2(file_path, dest)
        print(f" -> Copied {fname}")

print("\nALL SUPPORTED FORMATS SUCCESSFULLY GENERATED AND COPIED!")
